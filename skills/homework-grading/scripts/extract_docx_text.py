#!/usr/bin/env python
import argparse
import json
import sys
from pathlib import Path


def fail(message: str, code: int = 1) -> None:
    print(f"ERROR: {message}", file=sys.stderr)
    raise SystemExit(code)


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract ordered text, paragraphs, and tables from DOCX.")
    parser.add_argument("--input", required=True, help="Input DOCX path.")
    parser.add_argument("--output", required=True, help="Output JSON path.")
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output)
    if not input_path.exists():
        fail(f"Input DOCX not found: {input_path}")
    if input_path.suffix.lower() != ".docx":
        fail(f"Input must be a DOCX file: {input_path}")

    try:
        from docx import Document
    except ImportError:
        fail("Missing dependency python-docx. Install with: pip install python-docx")

    try:
        doc = Document(input_path)
        paragraphs = []
        ordered_parts = []
        for index, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text or ""
            paragraphs.append({"index": index, "text": text, "style": paragraph.style.name if paragraph.style else None})
            if text.strip():
                ordered_parts.append(text)

        tables = []
        for table_index, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append({"index": table_index, "rows": rows})
            for row in rows:
                ordered_parts.append(" | ".join(cell.strip() for cell in row))
    except Exception as exc:
        fail(f"Failed to extract DOCX text from {input_path}: {exc}")

    result = {
        "source_file": str(input_path),
        "paragraphs": paragraphs,
        "tables": tables,
        "full_text": "\n".join(ordered_parts),
        "extraction_method": "python-docx",
    }
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


if __name__ == "__main__":
    main()
